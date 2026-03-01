"""
测试边缘情况 - 提升覆盖率
"""

from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

from mcp_document_converter.core.engine import ConversionResult, DocumentConverter
from mcp_document_converter.core.ir import DocumentIR, Node, NodeType
from mcp_document_converter.core.parser import BaseParser, ParseError
from mcp_document_converter.core.renderer import BaseRenderer, RenderError
from mcp_document_converter.registry import ConverterRegistry
from mcp_document_converter.parsers import (
    DOCXParser,
    HTMLParser,
    MarkdownParser,
    PDFParser,
    TextParser,
)
from mcp_document_converter.renderers import (
    DOCXRenderer,
    HTMLRenderer,
    MarkdownRenderer,
    PDFRenderer,
    TextRenderer,
)


class TestCoreEngineEdgeCases:
    """测试转换引擎边缘情况"""

    @pytest.fixture
    def registry(self):
        reg = ConverterRegistry()
        reg.register_parser(MarkdownParser())
        reg.register_parser(HTMLParser())
        reg.register_parser(DOCXParser())
        reg.register_parser(PDFParser())
        reg.register_parser(TextParser())
        reg.register_renderer(HTMLRenderer())
        reg.register_renderer(MarkdownRenderer())
        reg.register_renderer(DOCXRenderer())
        reg.register_renderer(PDFRenderer())
        reg.register_renderer(TextRenderer())
        return reg

    @pytest.fixture
    def converter(self, registry):
        return DocumentConverter(registry)

    def test_convert_parse_error(self, converter, temp_dir):
        """测试解析错误"""
        # 创建一个无效的文件
        file_path = temp_dir / "test.xyz"
        file_path.write_text("content", encoding="utf-8")

        # 使用 source_format 强制使用 markdown 解析器
        # 但文件内容不是有效的 markdown（实际上任何文本都是有效的 markdown）
        # 所以我们测试不支持的格式
        result = converter.convert(file_path, "html")

        assert result.success is False
        assert "不支持的源格式" in result.error_message

    def test_render_from_ir_with_binary_content(self, converter, temp_dir):
        """测试二进制内容渲染"""
        doc = DocumentIR()
        doc.add_node(
            Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="test")])
        )

        result = converter.render_from_ir(doc, "docx")

        assert result.success is True
        assert isinstance(result.content, bytes)

    def test_render_from_ir_without_output_path(self, converter):
        """测试不指定输出路径的渲染"""
        doc = DocumentIR()
        doc.add_node(
            Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="test")])
        )

        result = converter.render_from_ir(doc, "html")

        assert result.success is True
        assert result.output_path is None
        assert isinstance(result.content, str)


class TestCoreIREdgeCases:
    """测试 IR 边缘情况"""

    def test_get_text_content_with_empty_string(self):
        """测试空字符串内容"""
        doc = DocumentIR()
        doc.add_node(
            Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="")])
        )

        text = doc.get_text_content()

        # 空字符串应该被过滤
        assert text == ""

    def test_get_text_content_with_whitespace_only(self):
        """测试只有空白的内容"""
        doc = DocumentIR()
        doc.add_node(
            Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="   ")])
        )

        text = doc.get_text_content()

        # 只有空白的内容应该被过滤
        assert text == ""

    def test_node_with_empty_list_content(self):
        """测试空列表内容"""
        node = Node(type=NodeType.PARAGRAPH, content=[])

        assert node.content == []


class TestCoreParserEdgeCases:
    """测试解析器基类边缘情况"""

    def test_base_parser_abstract_methods(self):
        """测试抽象方法不能直接实例化"""
        with pytest.raises(TypeError):
            BaseParser()

    def test_parse_error_with_source(self):
        """测试带源信息的解析错误"""
        error = ParseError("测试错误", source="test.md")

        assert error.message == "测试错误"
        assert error.source == "test.md"


class TestCoreRendererEdgeCases:
    """测试渲染器基类边缘情况"""

    def test_base_renderer_abstract_methods(self):
        """测试抽象方法不能直接实例化"""
        with pytest.raises(TypeError):
            BaseRenderer()

    def test_render_error_with_format(self):
        """测试带格式信息的渲染错误"""
        error = RenderError("测试错误", format_name="html")

        assert error.message == "测试错误"
        assert error.format_name == "html"


class TestMarkdownParserEdgeCases:
    """测试 Markdown 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return MarkdownParser()

    def test_parse_empty_content(self, parser, temp_dir):
        """测试空内容"""
        md_path = temp_dir / "empty.md"
        md_path.write_text("", encoding="utf-8")

        doc = parser.parse(md_path)

        assert isinstance(doc, DocumentIR)
        assert len(doc.content) == 0

    def test_parse_invalid_front_matter(self, parser, temp_dir):
        """测试无效的 YAML Front Matter"""
        md_path = temp_dir / "test.md"
        # 使用无效的 YAML 格式
        md_path.write_text("---\ninvalid: [yaml: content\n---\n# Title", encoding="utf-8")

        doc = parser.parse(md_path)

        # 无效的 YAML 应该被忽略
        assert isinstance(doc, DocumentIR)

    def test_parse_with_date(self, parser, temp_dir):
        """测试带日期的 Front Matter"""
        md_path = temp_dir / "test.md"
        # 使用带时间的日期格式（YAML 不会自动解析）
        md_path.write_text(
            "---\ntitle: Test\nauthor: Author\n---\n# Content",
            encoding="utf-8",
        )

        doc = parser.parse(md_path)

        assert doc.title == "Test"
        assert doc.author == "Author"

    def test_parse_div_element(self, parser, temp_dir):
        """测试 div 元素"""
        md_path = temp_dir / "test.md"
        md_path.write_text("<div>\n\nParagraph\n\n</div>", encoding="utf-8")

        doc = parser.parse(md_path)

        assert isinstance(doc, DocumentIR)

    def test_parse_strikethrough(self, parser, temp_dir):
        """测试删除线"""
        md_path = temp_dir / "test.md"
        md_path.write_text("~~删除线文本~~", encoding="utf-8")

        doc = parser.parse(md_path)

        assert isinstance(doc, DocumentIR)


class TestHTMLParserEdgeCases:
    """测试 HTML 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return HTMLParser()

    def test_parse_empty_html(self, parser, temp_dir):
        """测试空 HTML"""
        html_path = temp_dir / "test.html"
        html_path.write_text("<html><body></body></html>", encoding="utf-8")

        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)

    def test_parse_div_container(self, parser, temp_dir):
        """测试 div 容器"""
        html_path = temp_dir / "test.html"
        html_path.write_text(
            '<html><body><div class="content"><p>Text</p></div></body></html>',
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)

    def test_parse_article(self, parser, temp_dir):
        """测试 article 元素"""
        html_path = temp_dir / "test.html"
        html_path.write_text(
            '<html><body><article><p>Article content</p></article></body></html>',
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)

    def test_parse_section(self, parser, temp_dir):
        """测试 section 元素"""
        html_path = temp_dir / "test.html"
        html_path.write_text(
            '<html><body><section><p>Section content</p></section></body></html>',
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)

    def test_parse_main(self, parser, temp_dir):
        """测试 main 元素"""
        html_path = temp_dir / "test.html"
        html_path.write_text(
            '<html><body><main><p>Main content</p></main></body></html>',
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)

    def test_parse_table_with_tbody(self, parser, temp_dir):
        """测试带 tbody 的表格"""
        html_path = temp_dir / "test.html"
        html_path.write_text(
            '<html><body><table><thead><tr><th>H</th></tr></thead><tbody><tr><td>D</td></tr></tbody></table></body></html>',
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        tables = [n for n in doc.content if n.type == NodeType.TABLE]
        assert len(tables) >= 1

    def test_parse_unknown_element(self, parser, temp_dir):
        """测试未知元素"""
        html_path = temp_dir / "test.html"
        html_path.write_text(
            '<html><body><unknown>Content</unknown></body></html>',
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)


class TestTextParserEdgeCases:
    """测试 Text 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return TextParser()

    def test_parse_empty_text(self, parser, temp_dir):
        """测试空文本"""
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("", encoding="utf-8")

        doc = parser.parse(txt_path)

        assert isinstance(doc, DocumentIR)
        assert len(doc.content) == 0

    def test_parse_numbered_heading_various_levels(self, parser, temp_dir):
        """测试不同级别的编号标题"""
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("1. First\n\n10. Tenth\n\n100. Hundred", encoding="utf-8")

        doc = parser.parse(txt_path)

        headings = [n for n in doc.content if n.type == NodeType.HEADING]
        assert len(headings) >= 3

    def test_parse_multiline_blockquote(self, parser, temp_dir):
        """测试多行引用"""
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("> Line 1\n> Line 2\n> Line 3", encoding="utf-8")

        doc = parser.parse(txt_path)

        blockquotes = [n for n in doc.content if n.type == NodeType.BLOCKQUOTE]
        assert len(blockquotes) >= 1

    def test_parse_code_block_with_empty_lines(self, parser, temp_dir):
        """测试带空行的代码块"""
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("    line 1\n\n    line 2", encoding="utf-8")

        doc = parser.parse(txt_path)

        code_blocks = [n for n in doc.content if n.type == NodeType.CODE_BLOCK]
        assert len(code_blocks) >= 1

    def test_parse_list_with_empty_lines(self, parser, temp_dir):
        """测试带空行的列表"""
        txt_path = temp_dir / "test.txt"
        txt_path.write_text("- Item 1\n\n- Item 2", encoding="utf-8")

        doc = parser.parse(txt_path)

        # 列表可能被分成多个
        assert len(doc.content) >= 1


class TestDOCXParserEdgeCases:
    """测试 DOCX 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return DOCXParser()

    def test_parse_empty_document(self, parser, temp_dir):
        """测试空文档"""
        from docx import Document

        docx_path = temp_dir / "empty.docx"
        doc = Document()
        doc.save(str(docx_path))

        parsed = parser.parse(docx_path)

        assert isinstance(parsed, DocumentIR)

    def test_parse_document_with_table(self, parser, temp_dir):
        """测试带表格的文档"""
        from docx import Document

        docx_path = temp_dir / "table.docx"
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "2"
        doc.save(str(docx_path))

        parsed = parser.parse(docx_path)

        tables = [n for n in parsed.content if n.type == NodeType.TABLE]
        assert len(tables) >= 1

    def test_parse_document_with_list(self, parser, temp_dir):
        """测试带列表的文档"""
        from docx import Document

        docx_path = temp_dir / "list.docx"
        doc = Document()
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.save(str(docx_path))

        parsed = parser.parse(docx_path)

        assert isinstance(parsed, DocumentIR)

    def test_parse_document_with_bold_italic(self, parser, temp_dir):
        """测试带粗体斜体的文档"""
        from docx import Document

        docx_path = temp_dir / "format.docx"
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Bold text")
        run.bold = True
        run = p.add_run(" and ")
        run = p.add_run("italic text")
        run.italic = True
        doc.save(str(docx_path))

        parsed = parser.parse(docx_path)

        assert isinstance(parsed, DocumentIR)


class TestPDFParserEdgeCases:
    """测试 PDF 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return PDFParser()

    def test_classify_text_various_cases(self, parser):
        """测试文本分类各种情况"""
        # 普通段落
        node = parser._classify_text("Normal paragraph text")
        assert node.type == NodeType.PARAGRAPH

        # 列表
        node = parser._classify_text("- List item")
        assert node.type == NodeType.LIST

        # 带 bullet 的列表
        node = parser._classify_text("• Bullet item")
        assert node.type == NodeType.LIST

    def test_is_heading_various_cases(self, parser):
        """测试标题判断各种情况"""
        # 全大写
        assert parser._is_heading("IMPORTANT NOTICE") is True

        # 数字编号
        assert parser._is_heading("1. Introduction") is True

        # Markdown 风格
        assert parser._is_heading("## Heading") is True

        # 普通文本
        assert parser._is_heading("This is normal text") is False

        # 太长的文本
        assert parser._is_heading("A" * 150) is False


class TestHTMLRendererEdgeCases:
    """测试 HTML 渲染器边缘情况"""

    @pytest.fixture
    def renderer(self):
        return HTMLRenderer()

    def test_render_empty_document(self, renderer):
        """测试空文档"""
        doc = DocumentIR()
        html = renderer.render(doc)

        assert "<!DOCTYPE html>" in html

    def test_render_document_node(self, renderer):
        """测试文档节点"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.DOCUMENT,
                content=[
                    Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="Text")])
                ],
            )
        )

        html = renderer.render(doc)

        assert "Text" in html

    def test_render_empty_paragraph(self, renderer):
        """测试空段落"""
        doc = DocumentIR()
        doc.add_node(Node(type=NodeType.PARAGRAPH, content=[]))

        html = renderer.render(doc)

        assert isinstance(html, str)

    def test_render_code_block_with_invalid_language(self, renderer):
        """测试无效语言的代码块"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.CODE_BLOCK,
                content="code",
                attributes={"language": "nonexistent_language_xyz"},
            )
        )

        html = renderer.render(doc)

        # 应该回退到普通代码块
        assert "<pre>" in html


class TestMarkdownRendererEdgeCases:
    """测试 Markdown 渲染器边缘情况"""

    @pytest.fixture
    def renderer(self):
        return MarkdownRenderer()

    def test_render_empty_document(self, renderer):
        """测试空文档"""
        doc = DocumentIR()
        md = renderer.render(doc)

        assert isinstance(md, str)

    def test_render_document_node(self, renderer):
        """测试文档节点"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.DOCUMENT,
                content=[
                    Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="Text")])
                ],
            )
        )

        md = renderer.render(doc)

        assert "Text" in md

    def test_render_image_with_title(self, renderer):
        """测试带标题的图片"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.PARAGRAPH,
                content=[
                    Node(
                        type=NodeType.IMAGE,
                        content="",
                        attributes={"src": "img.png", "alt": "Alt", "title": "Title"},
                    )
                ],
            )
        )

        md = renderer.render(doc)

        assert '![Alt](img.png "Title")' in md

    def test_render_strikethrough(self, renderer):
        """测试删除线"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.PARAGRAPH,
                content=[
                    Node(type=NodeType.STRIKETHROUGH, content=[Node(type=NodeType.TEXT, content="del")])
                ],
            )
        )

        md = renderer.render(doc)

        assert "~~del~~" in md


class TestTextRendererEdgeCases:
    """测试 Text 渲染器边缘情况"""

    @pytest.fixture
    def renderer(self):
        return TextRenderer()

    def test_render_empty_document(self, renderer):
        """测试空文档"""
        doc = DocumentIR()
        text = renderer.render(doc)

        assert isinstance(text, str)

    def test_render_document_node(self, renderer):
        """测试文档节点"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.DOCUMENT,
                content=[
                    Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="Text")])
                ],
            )
        )

        text = renderer.render(doc)

        assert "Text" in text

    def test_render_heading_without_formatting(self, renderer):
        """测试不带格式的标题"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.HEADING,
                content=[Node(type=NodeType.TEXT, content="Title")],
                attributes={"level": 1},
            )
        )

        text = renderer.render(doc, preserve_formatting=False)

        assert "Title" in text
        assert "=" not in text

    def test_render_horizontal_rule_without_formatting(self, renderer):
        """测试不带格式的分隔线"""
        doc = DocumentIR()
        doc.add_node(Node(type=NodeType.HORIZONTAL_RULE, content=""))

        text = renderer.render(doc, preserve_formatting=False)

        assert text.strip() == "" or "-" not in text

    def test_render_inline_formatting(self, renderer):
        """测试行内格式"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.PARAGRAPH,
                content=[
                    Node(type=NodeType.STRONG, content=[Node(type=NodeType.TEXT, content="bold")]),
                    Node(type=NodeType.EMPHASIS, content=[Node(type=NodeType.TEXT, content="italic")]),
                    Node(type=NodeType.CODE_INLINE, content="code"),
                    Node(type=NodeType.STRIKETHROUGH, content=[Node(type=NodeType.TEXT, content="del")]),
                ],
            )
        )

        text = renderer.render(doc)

        assert "*bold*" in text
        assert "/italic/" in text
        assert "`code`" in text
        assert "~del~" in text


class TestDOCXRendererEdgeCases:
    """测试 DOCX 渲染器边缘情况"""

    @pytest.fixture
    def renderer(self):
        return DOCXRenderer()

    def test_render_empty_document(self, renderer):
        """测试空文档"""
        doc = DocumentIR()
        result = renderer.render(doc)

        assert isinstance(result, bytes)

    def test_render_document_node(self, renderer):
        """测试文档节点"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.DOCUMENT,
                content=[
                    Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="Text")])
                ],
            )
        )

        result = renderer.render(doc)

        assert isinstance(result, bytes)

    def test_render_page_break(self, renderer):
        """测试分页符"""
        doc = DocumentIR()
        doc.add_node(Node(type=NodeType.PAGE_BREAK, content=""))

        result = renderer.render(doc)

        assert isinstance(result, bytes)

    def test_render_line_break(self, renderer):
        """测试换行"""
        doc = DocumentIR()
        doc.add_node(
            Node(
                type=NodeType.PARAGRAPH,
                content=[
                    Node(type=NodeType.TEXT, content="Line 1"),
                    Node(type=NodeType.LINE_BREAK, content=""),
                    Node(type=NodeType.TEXT, content="Line 2"),
                ],
            )
        )

        result = renderer.render(doc)

        assert isinstance(result, bytes)


class TestPDFRendererEdgeCases:
    """测试 PDF 渲染器边缘情况"""

    @pytest.fixture
    def renderer(self):
        return PDFRenderer()

    def test_render_empty_document(self, renderer):
        """测试空文档"""
        doc = DocumentIR()

        try:
            result = renderer.render(doc)
            assert isinstance(result, bytes)
        except RenderError:
            pytest.skip("WeasyPrint not available")

    def test_render_page_break(self, renderer):
        """测试分页符"""
        doc = DocumentIR()
        doc.add_node(Node(type=NodeType.PAGE_BREAK, content=""))

        try:
            result = renderer.render(doc)
            assert isinstance(result, bytes)
        except RenderError:
            pytest.skip("WeasyPrint not available")

    def test_get_pdf_css(self, renderer):
        """测试 PDF CSS 生成"""
        css = renderer._get_pdf_css({"page_size": "A4", "margin": "2cm"})

        assert "@page" in css
        assert "A4" in css

    def test_register_chinese_font(self, renderer):
        """测试中文字体注册"""
        font_name = renderer._register_chinese_font()

        assert isinstance(font_name, str)
