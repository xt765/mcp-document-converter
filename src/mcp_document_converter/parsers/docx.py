"""
DOCX 解析器 - 将 Word 文档解析为中间表示
"""

from pathlib import Path
from typing import Any, List, Optional, Union

from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from ..core.ir import DocumentIR, Node, NodeType
from ..core.parser import BaseParser, ParseError


class DOCXParser(BaseParser):
    """
    Word 文档解析器
    
    将 .docx 文件解析为中间表示。
    """
    
    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]
    
    @property
    def format_name(self) -> str:
        return "docx"
    
    @property
    def mime_types(self) -> List[str]:
        return ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    
    def parse(self, source: Union[str, Path, bytes], **options: Any) -> DocumentIR:
        """
        解析 DOCX 文档
        
        Args:
            source: DOCX 文件路径或二进制内容
            **options: 解析选项
        
        Returns:
            DocumentIR: 文档的中间表示
        """
        try:
            # 读取文档
            if isinstance(source, (str, Path)):
                doc = DocxDocument(str(source))
            else:
                from io import BytesIO
                doc = DocxDocument(BytesIO(source))
            
            document = DocumentIR()
            
            # 提取文档属性
            core_props = doc.core_properties
            document.title = core_props.title
            document.author = core_props.author
            if core_props.created:
                document.created_at = core_props.created
            if core_props.modified:
                document.modified_at = core_props.modified
            
            # 提取其他元数据
            document.metadata = {
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
                "comments": core_props.comments,
            }
            
            # 解析文档内容
            for element in doc.element.body:
                node = self._parse_element(element, doc)
                if node:
                    document.add_node(node)
            
            return document
            
        except Exception as e:
            raise ParseError(f"DOCX 解析失败: {str(e)}")
    
    def _parse_element(self, element, doc: DocxDocument) -> Optional[Node]:
        """解析文档元素为 IR 节点"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn
        
        tag = element.tag
        
        # 段落
        if tag.endswith('p'):
            paragraph = next((p for p in doc.paragraphs if p._element is element), None)
            if paragraph:
                return self._parse_paragraph(paragraph)
        
        # 表格
        elif tag.endswith('tbl'):
            table = next((t for t in doc.tables if t._element is element), None)
            if table:
                return self._parse_table(table)
        
        return None
    
    def _parse_paragraph(self, paragraph: DocxParagraph) -> Optional[Node]:
        """解析段落"""
        # 检查是否为标题
        style_name = paragraph.style.name if paragraph.style else ""
        
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
            
            return Node(
                type=NodeType.HEADING,
                content=self._parse_runs(paragraph),
                attributes={"level": level}
            )
        
        # 检查是否为列表项
        elif paragraph._p.pPr is not None:
            numPr = paragraph._p.pPr.numPr
            if numPr is not None:
                # 这是一个列表项
                return Node(
                    type=NodeType.LIST_ITEM,
                    content=self._parse_runs(paragraph)
                )
        
        # 普通段落
        content = self._parse_runs(paragraph)
        if content:
            return Node(
                type=NodeType.PARAGRAPH,
                content=content
            )
        
        return None
    
    def _parse_runs(self, paragraph: DocxParagraph) -> List[Node]:
        """解析段落中的 runs（行内内容）"""
        nodes = []
        
        for run in paragraph.runs:
            text = run.text
            if not text.strip():
                continue
            
            # 检查样式
            is_bold = run.bold or (run.font and run.font.bold)
            is_italic = run.italic or (run.font and run.font.italic)
            is_code = run.font and run.font.name in ['Courier New', 'Consolas', 'Monaco']
            
            node = Node(type=NodeType.TEXT, content=text)
            
            if is_bold and is_italic:
                node = Node(type=NodeType.STRONG, content=[
                    Node(type=NodeType.EMPHASIS, content=[node])
                ])
            elif is_bold:
                node = Node(type=NodeType.STRONG, content=[node])
            elif is_italic:
                node = Node(type=NodeType.EMPHASIS, content=[node])
            
            if is_code:
                node = Node(type=NodeType.CODE_INLINE, content=text)
            
            nodes.append(node)
        
        return nodes
    
    def _parse_table(self, table: DocxTable) -> Node:
        """解析表格"""
        rows = []
        
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_content = []
                for paragraph in cell.paragraphs:
                    node = self._parse_paragraph(paragraph)
                    if node:
                        cell_content.append(node)
                
                cells.append(Node(
                    type=NodeType.TABLE_CELL,
                    content=cell_content,
                    attributes={"header": row_idx == 0}
                ))
            
            rows.append(Node(type=NodeType.TABLE_ROW, content=cells))
        
        return Node(type=NodeType.TABLE, content=rows)
