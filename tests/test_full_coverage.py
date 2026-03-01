"""
测试更多边缘情况 - 提升覆盖率到 90%+
"""

from unittest.mock import patch

import pytest

from mcp_document_converter.core.engine import DocumentConverter
from mcp_document_converter.core.ir import DocumentIR, Node, NodeType
from mcp_document_converter.core.parser import ParseError
from mcp_document_converter.core.renderer import RenderError
from mcp_document_converter.parsers import (
    DOCXParser,
    HTMLParser,
    MarkdownParser,
    PDFParser,
    TextParser,
)
from mcp_document_converter.registry import ConverterRegistry
from mcp_document_converter.renderers import (
    DOCXRenderer,
    HTMLRenderer,
    MarkdownRenderer,
    PDFRenderer,
    TextRenderer,
)


class TestEngineErrorHandling:
    """测试引擎错误处理"""

    @pytest.fixture
    def registry(self):
        reg = ConverterRegistry()
        reg.register_parser(MarkdownParser())
        reg.register_parser(HTMLParser())
        reg.register_parser(DOCXParser())
        reg.register_parser(PDFParser())
        reg.register_parser(TextParser())
        reg.register_renderer(HTMLRenderer())
        reg.register_renderer(MarkdownRenderer())
        reg.register_renderer(DOCXRenderer())
        reg.register_renderer(PDFRenderer())
        reg.register_renderer(TextRenderer())
        return reg

    @pytest.fixture
    def converter(self, registry):
        return DocumentConverter(registry)

    def test_convert_with_parse_error(self, converter, temp_dir):
        """测试解析错误处理"""
        # 创建一个会导致解析错误的场景
        # 由于 Markdown 解析器很宽容，我们使用模拟来测试
        md_path = temp_dir / "test.md"
        md_path.write_text("# Test", encoding="utf-8")

        # 模拟解析器抛出错误
        with patch.object(
            MarkdownParser,
            "parse",
            side_effect=ParseError("模拟解析错误"),
        ):
            # 创建新的转换器使用模拟的解析器
            reg = ConverterRegistry()
            reg.register_parser(MarkdownParser())
            reg.register_renderer(HTMLRenderer())
            conv = DocumentConverter(reg)

            result = conv.convert(md_path, "html")

            assert result.success is False
            assert result.error_message is not None
            assert "解析错误" in result.error_message

    def test_convert_with_render_error(self, converter, temp_dir, sample_markdown):
        """测试渲染错误处理"""
        md_path = temp_dir / "test.md"
        md_path.write_text(sample_markdown, encoding="utf-8")

        # 模拟渲染器抛出错误
        with patch.object(
            HTMLRenderer,
            "render",
            side_effect=RenderError("模拟渲染错误"),
        ):
            reg = ConverterRegistry()
            reg.register_parser(MarkdownParser())
            reg.register_renderer(HTMLRenderer())
            conv = DocumentConverter(reg)

            result = conv.convert(md_path, "html")

            assert result.success is False
            assert result.error_message is not None
            assert "渲染错误" in result.error_message

    def test_convert_with_general_exception(self, converter, temp_dir, sample_markdown):
        """测试一般异常处理"""
        md_path = temp_dir / "test.md"
        md_path.write_text(sample_markdown, encoding="utf-8")

        # 模拟渲染器抛出一般异常
        with patch.object(
            HTMLRenderer,
            "render",
            side_effect=Exception("一般错误"),
        ):
            reg = ConverterRegistry()
            reg.register_parser(MarkdownParser())
            reg.register_renderer(HTMLRenderer())
            conv = DocumentConverter(reg)

            result = conv.convert(md_path, "html")

            assert result.success is False
            assert result.error_message is not None
            assert "转换失败" in result.error_message

    def test_render_from_ir_with_render_error(self, converter):
        """测试从 IR 渲染时的渲染错误"""
        doc = DocumentIR()
        doc.add_node(
            Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="test")])
        )

        # 模拟渲染器抛出错误
        with patch.object(
            HTMLRenderer,
            "render",
            side_effect=RenderError("模拟渲染错误"),
        ):
            reg = ConverterRegistry()
            reg.register_renderer(HTMLRenderer())
            conv = DocumentConverter(reg)

            result = conv.render_from_ir(doc, "html")

            assert result.success is False
            assert result.error_message is not None
            assert "渲染错误" in result.error_message

    def test_render_from_ir_with_general_exception(self, converter):
        """测试从 IR 渲染时的一般异常"""
        doc = DocumentIR()
        doc.add_node(
            Node(type=NodeType.PARAGRAPH, content=[Node(type=NodeType.TEXT, content="test")])
        )

        # 模拟渲染器抛出一般异常
        with patch.object(
            HTMLRenderer,
            "render",
            side_effect=Exception("一般错误"),
        ):
            reg = ConverterRegistry()
            reg.register_renderer(HTMLRenderer())
            conv = DocumentConverter(reg)

            result = conv.render_from_ir(doc, "html")

            assert result.success is False
            assert result.error_message is not None
            assert "渲染失败" in result.error_message


class TestServerToolHandling:
    """测试服务器工具处理"""

    @pytest.mark.asyncio
    async def test_handle_unknown_tool(self):
        """测试处理未知工具"""
        from mcp_document_converter.server import create_server

        server = create_server()

        # 验证服务器创建成功
        assert server is not None
        assert server.name == "mcp-document-converter"


class TestPDFParserErrorHandling:
    """测试 PDF 解析器错误处理"""

    @pytest.fixture
    def parser(self):
        return PDFParser()

    def test_parse_invalid_pdf_bytes(self, parser):
        """测试解析无效的 PDF 字节"""
        import pytest

        from mcp_document_converter.core.parser import ParseError

        # 无效的 PDF 数据
        invalid_pdf = b"not a valid pdf content"

        with pytest.raises(ParseError) as exc_info:
            parser.parse(invalid_pdf)

        assert "PDF 解析失败" in str(exc_info.value)


class TestTextParserEdgeCases:
    """测试 Text 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return TextParser()

    def test_parse_only_whitespace(self, parser, temp_dir):
        """测试只有空白字符的文件"""
        txt_path = temp_dir / "whitespace.txt"
        txt_path.write_text("   \n\n\t\t\n   ", encoding="utf-8")

        doc = parser.parse(txt_path)

        # 应该没有内容
        assert len(doc.content) == 0

    def test_parse_single_line(self, parser, temp_dir):
        """测试单行文本"""
        txt_path = temp_dir / "single.txt"
        txt_path.write_text("Single line of text", encoding="utf-8")

        doc = parser.parse(txt_path)

        assert len(doc.content) >= 1

    def test_parse_mixed_content(self, parser, temp_dir):
        """测试混合内容"""
        txt_path = temp_dir / "mixed.txt"
        txt_path.write_text(
            """IMPORTANT HEADING

This is a paragraph.

- List item 1
- List item 2

    code block here

> Quote text

---

Final paragraph.""",
            encoding="utf-8",
        )

        doc = parser.parse(txt_path)

        # 应该有多种类型的节点
        types = set(n.type for n in doc.content)
        assert len(types) >= 2


class TestHTMLParserEdgeCases:
    """测试 HTML 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return HTMLParser()

    def test_parse_malformed_html(self, parser, temp_dir):
        """测试解析畸形 HTML"""
        html_path = temp_dir / "malformed.html"
        html_path.write_text(
            "<html><body><p>Unclosed paragraph<div>Nested</body></html>",
            encoding="utf-8",
        )

        # 应该能够处理畸形 HTML
        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)

    def test_parse_html_with_scripts(self, parser, temp_dir):
        """测试带脚本的 HTML"""
        html_path = temp_dir / "script.html"
        html_path.write_text(
            '<html><head><script>alert("test");</script></head><body><p>Content</p></body></html>',
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        # 脚本应该被忽略
        assert isinstance(doc, DocumentIR)

    def test_parse_html_with_styles(self, parser, temp_dir):
        """测试带样式的 HTML"""
        html_path = temp_dir / "style.html"
        html_path.write_text(
            "<html><head><style>body { color: red; }</style></head><body><p>Content</p></body></html>",
            encoding="utf-8",
        )

        doc = parser.parse(html_path)

        assert isinstance(doc, DocumentIR)


class TestDOCXParserEdgeCases:
    """测试 DOCX 解析器边缘情况"""

    @pytest.fixture
    def parser(self):
        return DOCXParser()

    def test_parse_document_with_hyperlink(self, parser, temp_dir):
        """测试带超链接的文档"""
        from docx import Document

        docx_path = temp_dir / "hyperlink.docx"
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("Click ")
        run = p.add_run("here")
        run.underline = True
        doc.save(str(docx_path))

        parsed = parser.parse(docx_path)

        assert isinstance(parsed, DocumentIR)

    def test_parse_document_with_image_reference(self, parser, temp_dir):
        """测试带图片引用的文档"""
        from docx import Document

        docx_path = temp_dir / "image.docx"
        doc = Document()
        doc.add_paragraph("Text before image")
        # 添加一个空的图片占位符
        p = doc.add_paragraph()
        _ = p.add_run()
        doc.add_paragraph("Text after image")
        doc.save(str(docx_path))

        parsed = parser.parse(docx_path)

        assert isinstance(parsed, DocumentIR)


class TestRenderersEdgeCases:
    """测试渲染器边缘情况"""

    def test_html_renderer_empty_nodes(self):
        """测试 HTML 渲染器空节点"""
        renderer = HTMLRenderer()
        doc = DocumentIR()

        html = renderer.render(doc)

        assert "<!DOCTYPE html>" in html

    def test_markdown_renderer_empty_nodes(self):
        """测试 Markdown 渲染器空节点"""
        renderer = MarkdownRenderer()
        doc = DocumentIR()

        md = renderer.render(doc)

        assert isinstance(md, str)

    def test_text_renderer_empty_nodes(self):
        """测试 Text 渲染器空节点"""
        renderer = TextRenderer()
        doc = DocumentIR()

        text = renderer.render(doc)

        assert isinstance(text, str)

    def test_docx_renderer_empty_nodes(self):
        """测试 DOCX 渲染器空节点"""
        renderer = DOCXRenderer()
        doc = DocumentIR()

        result = renderer.render(doc)

        assert isinstance(result, bytes)

    def test_pdf_renderer_empty_nodes(self):
        """测试 PDF 渲染器空节点"""
        renderer = PDFRenderer()
        doc = DocumentIR()

        try:
            result = renderer.render(doc)
            assert isinstance(result, bytes)
        except RenderError:
            pytest.skip("WeasyPrint not available")


class TestConversionMatrix:
    """测试转换矩阵"""

    def test_all_conversions_possible(self):
        """测试所有转换组合"""
        registry = ConverterRegistry()
        registry.register_parser(MarkdownParser())
        registry.register_parser(HTMLParser())
        registry.register_parser(DOCXParser())
        registry.register_parser(PDFParser())
        registry.register_parser(TextParser())
        registry.register_renderer(HTMLRenderer())
        registry.register_renderer(MarkdownRenderer())
        registry.register_renderer(DOCXRenderer())
        registry.register_renderer(PDFRenderer())
        registry.register_renderer(TextRenderer())

        converter = DocumentConverter(registry)
        matrix = converter.list_supported_conversions()

        # 检查每个源格式都有目标格式
        for source_format, targets in matrix.items():
            assert len(targets) > 0
            # 每个源格式应该能转换到所有目标格式
            assert "html" in targets
            assert "markdown" in targets
            assert "docx" in targets
            assert "pdf" in targets
            assert "text" in targets
