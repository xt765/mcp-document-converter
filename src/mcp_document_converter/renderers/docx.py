"""
DOCX 渲染器 - 将中间表示渲染为 Word 文档
"""

from typing import Any

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor

from ..core.ir import DocumentIR, Node, NodeType
from ..core.renderer import BaseRenderer, RenderError


class DOCXRenderer(BaseRenderer):
    """
    Word 文档渲染器

    将中间表示渲染为 .docx 格式。
    """

    @property
    def output_extension(self) -> str:
        return ".docx"

    @property
    def format_name(self) -> str:
        return "docx"

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def is_binary(self) -> bool:
        return True

    def render(self, document: DocumentIR, **options: Any) -> bytes:
        """
        将中间表示渲染为 DOCX

        Args:
            document: 文档的中间表示
            **options: 渲染选项

        Returns:
            DOCX 二进制数据
        """
        try:
            # 创建 Word 文档
            doc = DocxDocument()

            # 设置文档属性
            if document.title:
                doc.core_properties.title = document.title
            if document.author:
                doc.core_properties.author = document.author

            # 渲染内容
            for node in document.content:
                self._render_node(doc, node)

            # 保存到字节流
            from io import BytesIO

            output = BytesIO()
            doc.save(output)
            output.seek(0)

            return output.read()

        except Exception as e:
            raise RenderError(f"DOCX 渲染失败: {str(e)}")

    def _render_node(self, doc: DocxDocument, node: Node) -> None:
        """渲染单个节点"""
        if node.type == NodeType.HEADING:
            level = node.attributes.get("level", 1)
            content = self._extract_text(node)

            # 添加标题段落
            paragraph = doc.add_heading(content, level=level)

        elif node.type == NodeType.PARAGRAPH:
            paragraph = doc.add_paragraph()
            self._add_inline_content(paragraph, node)

        elif node.type == NodeType.CODE_BLOCK:
            code = node.content if isinstance(node.content, str) else ""
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(10)

            # 设置代码块样式
            paragraph.paragraph_format.left_indent = Inches(0.5)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)

        elif node.type == NodeType.LIST:
            self._render_list(doc, node)

        elif node.type == NodeType.TABLE:
            self._render_table(doc, node)

        elif node.type == NodeType.BLOCKQUOTE:
            # 创建引用段落
            content_nodes = node.content if isinstance(node.content, list) else []
            for child in content_nodes:
                if child.type == NodeType.PARAGRAPH:
                    paragraph = doc.add_paragraph()
                    self._add_inline_content(paragraph, child)
                    # 添加左边框效果（通过缩进模拟）
                    paragraph.paragraph_format.left_indent = Inches(0.3)
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)

        elif node.type == NodeType.HORIZONTAL_RULE:
            # 添加分隔线（使用下划线段落）
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("─" * 50)
            run.font.color.rgb = RGBColor(200, 200, 200)

        elif node.type == NodeType.DOCUMENT:
            # 递归渲染子节点
            for child in node.content if isinstance(node.content, list) else []:
                self._render_node(doc, child)

    def _render_list(self, doc: DocxDocument, node: Node) -> None:
        """渲染列表"""
        list_type = node.attributes.get("type", "unordered")
        items = node.content if isinstance(node.content, list) else []

        for i, item in enumerate(items):
            if not isinstance(item, Node):
                continue

            self._extract_text(item)

            if list_type == "ordered":
                paragraph = doc.add_paragraph(style="List Number")
            else:
                paragraph = doc.add_paragraph(style="List Bullet")

            self._add_inline_content(paragraph, item)

    def _render_table(self, doc: DocxDocument, node: Node) -> None:
        """渲染表格"""
        rows = node.content if isinstance(node.content, list) else []

        if not rows:
            return

        # 计算表格尺寸
        num_rows = len(rows)
        num_cols = max(len(row.content) if isinstance(row.content, list) else 0 for row in rows)

        if num_cols == 0:
            return

        # 创建表格
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"

        for row_idx, row in enumerate(rows):
            if not isinstance(row, Node) or row.type != NodeType.TABLE_ROW:
                continue

            cells = row.content if isinstance(row.content, list) else []

            for col_idx, cell in enumerate(cells):
                if col_idx >= num_cols:
                    break

                if not isinstance(cell, Node) or cell.type != NodeType.TABLE_CELL:
                    continue

                # 获取单元格
                docx_cell = table.rows[row_idx].cells[col_idx]

                # 清空默认段落
                docx_cell.text = ""
                paragraph = docx_cell.paragraphs[0]

                # 添加内容
                self._add_inline_content(paragraph, cell)

                # 设置表头样式
                is_header = cell.attributes.get("header", False) or row_idx == 0
                if is_header:
                    for run in paragraph.runs:
                        run.bold = True

    def _add_inline_content(self, paragraph, node: Node) -> None:
        """添加行内内容到段落"""
        if isinstance(node.content, str):
            paragraph.add_run(node.content)
            return

        if not isinstance(node.content, list):
            return

        for child in node.content:
            if isinstance(child, Node):
                if child.type == NodeType.TEXT:
                    paragraph.add_run(str(child.content))
                elif child.type == NodeType.STRONG:
                    run = paragraph.add_run(self._extract_text(child))
                    run.bold = True
                elif child.type == NodeType.EMPHASIS:
                    run = paragraph.add_run(self._extract_text(child))
                    run.italic = True
                elif child.type == NodeType.CODE_INLINE:
                    code = child.content if isinstance(child.content, str) else ""
                    run = paragraph.add_run(code)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                elif child.type == NodeType.LINK:
                    child.attributes.get("url", "")
                    text = self._extract_text(child)
                    run = paragraph.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                elif child.type == NodeType.STRIKETHROUGH:
                    run = paragraph.add_run(self._extract_text(child))
                    run.font.strike = True
                elif child.type == NodeType.LINE_BREAK:
                    paragraph.add_run().add_break()
                else:
                    self._add_inline_content(paragraph, child)

    def _extract_text(self, node: Node) -> str:
        """从节点中提取纯文本"""
        if isinstance(node.content, str):
            return node.content

        if not isinstance(node.content, list):
            return ""

        parts = []
        for child in node.content:
            if isinstance(child, Node):
                parts.append(self._extract_text(child))
            else:
                parts.append(str(child))

        return "".join(parts)
